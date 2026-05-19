"""DOCX text extraction using python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def extract(file_path: Path) -> str:
    """Extract all paragraph text from a DOCX file."""
    doc = Document(str(file_path))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)
